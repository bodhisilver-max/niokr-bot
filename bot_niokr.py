import os, re, base64, io
import telebot
from telebot import types
from docx import Document
from docx.oxml.ns import qn
from PIL import Image

TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "")
PASSWORDS = {"niokr2026": "admin", "niokr2026v": "viewer"}
DOC_PATH, HTML_PATH = "niokr.docx", "niokr_app.html"

REPORT, DEF_YEAR, GREEN = (2026, 2), 2026, (4, 11)
MONTHS = {"январь":1,"февраль":2,"март":3,"апрель":4,"май":5,"июнь":6,"июль":7,
          "август":8,"сентябрь":9,"октябрь":10,"ноябрь":11,"декабрь":12}
IMG_STYLE = ("display:block;width:calc(100% - 24px);margin:12px auto;"
             "border-radius:12px;box-shadow:0 2px 8px rgba(0,0,0,.3)")

def month_run(r):
    m = re.match(r"([а-яё]+)\s*(\d{4})?", r.text.strip().lower())
    mo = MONTHS.get(m.group(1)) if m else None
    if not mo: return None
    yr = int(m.group(2)) if m.group(2) else DEF_YEAR
    return (yr, (mo - 1) // 3 + 1)

def run_status(r):
    if r.font.strike: return "перенесен"
    hc = r.font.highlight_color
    hv = int(hc) if hc is not None else -1
    if hv in GREEN: return "выполнен"
    return "план"

def cell_shaded(cell):
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None: return False
    shd = tcPr.find(qn('w:shd'))
    if shd is None: return False
    return (shd.get(qn('w:fill')) or "").upper() not in ("", "AUTO", "FFFFFF")

def row_excluded(row, text):
    strike = any(r.font.strike for p in row.cells[0].paragraphs for r in p.runs if r.text.strip())
    return "исключ" in text.lower() or (strike and cell_shaded(row.cells[0]))

def cell_image(cell, doc):
    for blip in cell._tc.iter(qn('a:blip')):
        rid = blip.get(qn('r:embed'))
        if not rid: continue
        part = doc.part.related_parts.get(rid)
        if not part: continue
        try:
            im = Image.open(io.BytesIO(part.blob)); im.thumbnail((1200, 1200))
            buf = io.BytesIO(); im.convert("RGB").save(buf, "JPEG", quality=80)
            return "data:image/jpeg;base64," + base64.b64encode(buf.getvalue()).decode()
        except Exception:
            return "data:image/png;base64," + base64.b64encode(part.blob).decode()
    return ""

def build_data(path):
    doc = Document(path)
    tab = max(doc.tables, key=lambda t: len(t.rows))
    themes, cur = [], None
    for row in tab.rows:
        tx = [c.text.strip() for c in row.cells]
        if not any(tx): continue
        if len(set(tx)) == 1:
            if row_excluded(row, tx[0]):
                if cur is not None:
                    cur["excluded"] = True
                    cur["note"] = (cur.get("note", "") + " " + tx[0]).strip()
                else:
                    num = re.match(r"\d+", tx[0])
                    themes.append(dict(number=num.group() if num else str(len(themes)+1),
                                       name=tx[0], stages=[], excluded=True, note=tx[0], img=""))
            continue
        if tx[0] in ("1","Наименование темы") or tx[1] in ("2","Этапы работ"): continue
        if tx[0] and tx[0] == tx[1]: continue
        if not tx[1]: continue
        if tx[0] and (cur is None or tx[0] != cur["name"]):
            num = re.match(r"\d+", tx[0])
            cur = dict(number=num.group() if num else str(len(themes)+1), name=tx[0],
                       stages=[], excluded=row_excluded(row, tx[0]),
                       note=" ".join(x for x in tx[4:] if x),
                       img=cell_image(row.cells[0], doc))
            themes.append(cur)
        if cur and tx[1] and not cur["excluded"]:
            ev = []
            for ci, cell in enumerate(row.cells[2:4]):
                for p in cell.paragraphs:
                    for r in p.runs:
                        q = month_run(r)
                        if q: ev.append((ci, q, run_status(r), r.text.strip()))
            cur["stages"].append(dict(name=tx[1], ev=ev, act=tx[7] if len(tx) > 7 else "",
                                      note=" ".join(x for x in tx[4:] if x)))
        if cur and cur["excluded"] and tx[1]:
            extra = " ".join(x for x in tx[3:] if x)
            if extra: cur["note"] = (cur.get("note", "") + " " + extra).strip()
    DATA = []
    for t in themes:
        stl = []
        for s in t["stages"]:
            allst = [st for (ci, q, st, raw) in s["ev"]]
            low = s["note"].lower()
            if any("выполнен" in x for x in allst) or "выполнен" in low: status = "done"
            elif any("перенес" in x for x in allst) or "перенес" in low: status = "moved"
            else: status = "plan"
            m = re.search(r"(акт|приказ|письмо|протокол|решение).*$", s["note"], re.I)
            doc_ = m.group(0).strip(" .;") if m else (s["note"] if status != "plan" and s["note"] else s["act"])
            stl.append(dict(name=s["name"], status=status,
                            plan=[raw for (ci, q, st, raw) in s["ev"] if ci == 0],
                            prop=[raw for (ci, q, st, raw) in s["ev"] if ci == 1],
                            doc=doc_))
        color = "red" if t["excluded"] else ("green" if stl and all(x["status"] == "done" for x in stl) else "")
        DATA.append(dict(n=t["number"], name=t["name"], color=color, stages=stl,
                         note=t.get("note", ""), img=t.get("img", "")))
    return DATA

def esc(s): return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def build_html(DATA):
    COL = {"done": "#e8f5e9", "moved": "#bdbdbd", "plan": "#ffffff"}
    ST = {"done": "✅ выполнен", "moved": "⬜ перенесен", "plan": "📅 план"}
    h = ['''<!DOCTYPE html><html><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<style>
body{margin:0;font-family:Segoe UI,Arial}
.top{background:#0d47a1;color:#fff;padding:12px;font-size:17px;position:sticky;top:0}
.top a{color:#fff;text-decoration:none;font-size:20px;margin-right:8px}
.row{display:block;padding:12px 14px;border-bottom:1px solid #ddd;color:#000;text-decoration:none}
.row:active{opacity:.6}
.sec{display:none}
.sec:target{display:block;position:fixed;top:0;left:0;right:0;bottom:0;background:#fff;overflow-y:auto;z-index:10}
.pass{padding:16px}
.frame{position:relative;margin:12px}
.frame img{width:100%;display:block;border-radius:12px;box-shadow:0 2px 8px rgba(0,0,0,.3)}
.overlay{position:absolute;top:8%;right:2.5%;width:45%;max-height:84%;overflow-y:auto;background:rgba(255,255,255,.93);border-radius:8px}
.overlay .row{padding:10px 12px;font-size:14px}
@media(max-width:600px){
  .overlay{position:static;width:auto;max-height:none;background:#fff;margin-top:8px}
  .overlay .row{padding:12px 14px;font-size:15px}
}
</style></head><body>
<div id="top" class="top">Темы НИОКР — 2 квартал 2026</div>''']
    for i, t in enumerate(DATA):
        bg = "#ffebee" if t["color"] == "red" else "#e8f5e9" if t["color"] == "green" else "#ffffff"
        h.append(f'<a class="row" style="background:{bg}" href="#t{i}"><b>{t["n"]}.</b> {"📷 " if t.get("img") else ""}{esc(t["name"])}</a>')
    for i, t in enumerate(DATA):
        if t["color"] == "red":
            h.append(f'<div class="sec" id="t{i}"><div class="top"><a href="#top">←</a>Тема {t["n"]}</div>'
                     + (f'<img src="{t["img"]}" style="{IMG_STYLE}">' if t.get("img") else "")
                     + f'<div class="pass"><h3>❌ Тема удалена (исключена)</h3>'
                       f'<p><b>Документ:</b> {esc(t["note"]) or "—"}</p></div></div>')
            continue
        h.append(f'<div class="sec" id="t{i}"><div class="top"><a href="#top">←</a>Тема {t["n"]}</div>')
        if t.get("img"):
            h.append(f'<div class="frame"><img src="{t["img"]}"><div class="overlay">')
            for j, s in enumerate(t["stages"]):
                h.append(f'<a class="row" style="background:{COL[s["status"]]}" href="#t{i}s{j}">{esc(s["name"])}</a>')
            h.append('</div></div>')
        else:
            for j, s in enumerate(t["stages"]):
                h.append(f'<a class="row" style="background:{COL[s["status"]]}" href="#t{i}s{j}">{esc(s["name"])}</a>')
        h.append('</div>')
        for j, s in enumerate(t["stages"]):
            extra = f'<p><b>Предлагаемый срок:</b> {esc(", ".join(s["prop"]))}</p>' if s["prop"] else ""
            h.append(f'<div class="sec" id="t{i}s{j}"><div class="top"><a href="#t{i}">←</a>Этап</div>'
                     f'<div class="pass"><h3>{esc(s["name"])}</h3>'
                     f'<p><b>Статус:</b> {ST[s["status"]]}</p>'
                     f'<p><b>План:</b> {esc(", ".join(s["plan"])) or "—"}</p>{extra}'
                     f'<p><b>Документ:</b> {esc(s["doc"]) or "—"}</p></div></div>')
    h.append("</body></html>")
    with open(HTML_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(h))

DATA = []
if os.path.exists(DOC_PATH):
    DATA = build_data(DOC_PATH); build_html(DATA)
print(f"📚 Тем: {len(DATA)} | этапов: {sum(len(t['stages']) for t in DATA)}")

bot = telebot.TeleBot(TOKEN)
ROLE = {}

def txt_sum():
    if not DATA: return "📭 Данных пока нет. Админ, пришли .docx"
    st = [s for t in DATA for s in t['stages']]
    done = sum(1 for s in st if s['status'] == 'done')
    moved = sum(1 for s in st if s['status'] == 'moved')
    excl = sum(1 for t in DATA if t['color'] == 'red')
    return (f"📊 НИОКР, 2 квартал 2026\n\n📌 Тем: {len(DATA)} (исключено: {excl})\n"
            f"📌 Этапов: {len(st)}\n✅ выполнено: {done}\n↗️ перенесено: {moved}\n"
            f"⏳ в работе: {len(st) - done - moved}")

def kb():
    k = types.InlineKeyboardMarkup()
    k.add(types.InlineKeyboardButton("📊 Сводка", callback_data="sum"),
          types.InlineKeyboardButton("📂 Открыть темы", callback_data="file"))
    return k

@bot.message_handler(commands=["start"])
def start(m):
    if m.chat.id in ROLE:
        bot.send_message(m.chat.id, "Уже открыто 🙂", reply_markup=kb()); return
    s = bot.send_message(m.chat.id, "🔒 Статус НИОКР. Введите пароль:")
    bot.register_next_step_handler(s, check)

def check(m):
    role = PASSWORDS.get(m.text)
    if role:
        ROLE[m.chat.id] = role
        msg = "✅ Доступ открыт." + (" Ты админ: пришли .docx для обновления." if role == "admin" else "")
        bot.send_message(m.chat.id, msg, reply_markup=kb())
    else:
        bot.send_message(m.chat.id, "❌ Неверный пароль.")

@bot.message_handler(content_types=["document"])
def on_doc(m):
    global DATA
    if ROLE.get(m.chat.id) != "admin":
        bot.send_message(m.chat.id, "⛔ Обновлять данные может только админ."); return
    if not m.document.file_name.endswith(".docx"):
        bot.send_message(m.chat.id, "Нужен файл .docx"); return
    pf = bot.get_file(m.document.file_id)
    with open(DOC_PATH, "wb") as f:
        f.write(bot.download_file(pf.file_path))
    DATA = build_data(DOC_PATH); build_html(DATA)
    bot.send_message(m.chat.id, "✅ Данные обновлены!\n" + txt_sum(), reply_markup=kb())

@bot.callback_query_handler(func=lambda c: True)
def cb(c):
    if c.message.chat.id not in ROLE:
        bot.answer_callback_query(c.id, "Сначала /start и пароль"); return
    if c.data == "sum":
        bot.send_message(c.message.chat.id, txt_sum(), reply_markup=kb())
    elif c.data == "file":
        if os.path.exists(HTML_PATH):
            with open(HTML_PATH, "rb") as f: bot.send_document(c.message.chat.id, f)
        else:
            bot.send_message(c.message.chat.id, "📭 Сначала загрузи .docx")
    bot.answer_callback_query(c.id)

# ---------- RENDER: webhook + сервер ----------
from flask import Flask, request as fl_request
WEBHOOK_URL = os.environ.get("RENDER_EXTERNAL_URL", "") + "/" + TOKEN
app = Flask(__name__)

@app.route("/" + TOKEN, methods=["POST"])
def webhook():
    if fl_request.headers.get("content-type") == "application/json":
        bot.process_new_updates([telebot.types.Update.de_json(fl_request.get_json())])
    return "ok"

bot.remove_webhook()
bot.set_webhook(WEBHOOK_URL)
print("🤖 Бот в строю! Webhook:", WEBHOOK_URL)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 8000)))
